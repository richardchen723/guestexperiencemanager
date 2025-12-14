#!/usr/bin/env python3
"""
Document parser for extracting text from Word and PDF documents.
"""

import os
import logging
from typing import Dict, Any, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_document(file_path: str, mime_type: str) -> Dict[str, Any]:
    """
    Parse document and extract text content.
    
    Args:
        file_path: Path to the document file
        mime_type: MIME type of the document (e.g., 'application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    Returns:
        {
            'text': str,  # Extracted text
            'page_count': int,  # For PDFs
            'word_count': int,
            'metadata': dict  # Document metadata if available
        }
    
    Raises:
        ValueError: If file type is not supported or parsing fails
    """
    file_path_obj = Path(file_path)
    
    if not file_path_obj.exists():
        raise ValueError(f"File not found: {file_path}")
    
    # Determine parser based on MIME type or file extension
    if mime_type == 'application/pdf' or file_path_obj.suffix.lower() == '.pdf':
        return _parse_pdf(file_path)
    elif (mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or 
          mime_type == 'application/msword' or
          file_path_obj.suffix.lower() in ['.docx', '.doc']):
        return _parse_word(file_path)
    else:
        raise ValueError(f"Unsupported file type: {mime_type}")


def _parse_pdf(file_path: str) -> Dict[str, Any]:
    """Parse PDF document and extract text."""
    try:
        import pdfplumber
        
        text_parts = []
        page_count = 0
        metadata = {}
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            
            # Extract metadata if available
            if pdf.metadata:
                metadata = {
                    'title': pdf.metadata.get('Title'),
                    'author': pdf.metadata.get('Author'),
                    'subject': pdf.metadata.get('Subject'),
                    'creator': pdf.metadata.get('Creator'),
                    'producer': pdf.metadata.get('Producer'),
                    'creation_date': str(pdf.metadata.get('CreationDate')) if pdf.metadata.get('CreationDate') else None,
                    'modification_date': str(pdf.metadata.get('ModDate')) if pdf.metadata.get('ModDate') else None,
                }
            
            # Extract text from each page
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        full_text = '\n\n'.join(text_parts)
        word_count = len(full_text.split()) if full_text else 0
        
        return {
            'text': full_text,
            'page_count': page_count,
            'word_count': word_count,
            'metadata': metadata
        }
        
    except ImportError:
        # Fallback to PyPDF2 if pdfplumber is not available
        try:
            import PyPDF2
            
            text_parts = []
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title'),
                        'author': pdf_reader.metadata.get('/Author'),
                        'subject': pdf_reader.metadata.get('/Subject'),
                        'creator': pdf_reader.metadata.get('/Creator'),
                        'producer': pdf_reader.metadata.get('/Producer'),
                    }
                
                # Extract text from each page
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            full_text = '\n\n'.join(text_parts)
            word_count = len(full_text.split()) if full_text else 0
            
            return {
                'text': full_text,
                'page_count': page_count,
                'word_count': word_count,
                'metadata': metadata
            }
        except ImportError:
            # If no PDF library is available, return empty text but don't fail
            logger.warning("PDF parsing library not available. Document uploaded but content not extracted. Please install pdfplumber or PyPDF2.")
            return {
                'text': '',
                'page_count': 0,
                'word_count': 0,
                'metadata': {}
            }
    except Exception as e:
        logger.error(f"Error parsing PDF {file_path}: {e}", exc_info=True)
        # Return empty text instead of raising error - document can still be uploaded
        return {
            'text': '',
            'page_count': 0,
            'word_count': 0,
            'metadata': {}
        }


def _parse_word(file_path: str) -> Dict[str, Any]:
    """Parse Word document (.docx) and extract text."""
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # Extract text from all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        full_text = '\n\n'.join(text_parts)
        word_count = len(full_text.split()) if full_text else 0
        
        # Extract metadata if available
        metadata = {}
        if doc.core_properties:
            props = doc.core_properties
            metadata = {
                'title': props.title,
                'author': props.author,
                'subject': props.subject,
                'keywords': props.keywords,
                'comments': props.comments,
                'created': str(props.created) if props.created else None,
                'modified': str(props.modified) if props.modified else None,
            }
        
        # Estimate page count (rough approximation: ~500 words per page)
        estimated_pages = max(1, word_count // 500) if word_count > 0 else 1
        
        return {
            'text': full_text,
            'page_count': estimated_pages,
            'word_count': word_count,
            'metadata': metadata
        }
        
    except ImportError:
        raise ValueError("Word document parsing library (python-docx) not available. Please install python-docx.")
    except Exception as e:
        logger.error(f"Error parsing Word document {file_path}: {e}", exc_info=True)
        raise ValueError(f"Failed to parse Word document: {str(e)}")

